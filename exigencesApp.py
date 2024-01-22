import re
from docx import Document
import openpyxl
import tkinter as tk
#from tkinter import ttk
#from tkinter import scrolledtext
from tkinter import filedialog as fd
import os
import configparser
#from tkinter import messagebox

FICHIER_INI = 'exigencesApp.ini'

class Exigences_App:
    def __init__(self, root):
        # Initialiser le parseur de configuration
        self.config = configparser.ConfigParser()

        # création du fichier INI s'il n'existe pas
        self.__createIniFileIfNotExist(self.config)

        # Chargement du fichier INI
        self.config.read(FICHIER_INI) 

        self.lire_configuration()

        #Initialisation de l'IHM
        self.root = root
        self.root.geometry("1350x700+0+0")
        self.root.title("Exigences")

        self.specsPath = tk.StringVar()
        self.rtfPath = tk.StringVar()

        bg_color = "#6fa8dc"
        title = tk.Label(self.root, text="Exigences", font=('times new roman', 30, 'bold'), pady=2, bd=2, bg="#6fa8dc", fg="Black", relief="solid")
        title.pack(fill=tk.X)

        # Frame pour la selection du document de specs
        F1 = tk.LabelFrame(self.root, text="Fichier de spécifications ...", font=('times new roman', 14, 'bold'), bd=2, fg="Black", bg="#6fa8dc")
        F1.place(x=0, y=60, relwidth=1, height=75)
        cname_lbl = tk.Label(F1, text="Chemin :", bg=bg_color, font=('times new roman', 12, 'normal'))
        cname_lbl.grid(row=0, column=0, padx=20, pady=5)
        cname_txt = tk.Entry(F1, width=80, textvariable=self.specsPath, font='arial 8', bd=2, relief='ridge')
        cname_txt.grid(row=0, column=1, pady=5, padx=10)
        cname_txt.config(state=tk.DISABLED)

        # Lier la fonction de vérification au changement de texte dans le champ
        self.specsPath.trace_add("write", self.verifier_chemin_fichiers)


        open_button = tk.Button(
            F1,
            text='Parcourir',
            command=self.select_files
        )
        open_button.grid(row=0, column=2, pady=5, padx=10)        

        # Frame pour la selection du RTF
        F2 = tk.LabelFrame(self.root, text="Fichier de RTF ...", font=('times new roman', 14, 'bold'), bd=2, fg="Black", bg="#6fa8dc")
        F2.place(x=0, y=150, relwidth=1, height=75)
        cname_lbl_XL = tk.Label(F2, text="Chemin :", bg=bg_color, font=('times new roman', 12, 'normal'))
        cname_lbl_XL.grid(row=0, column=0, padx=20, pady=5)
        cname_txt_XL = tk.Entry(F2, width=80, textvariable=self.rtfPath, font='arial 8', bd=2, relief='ridge')
        cname_txt_XL.grid(row=0, column=1, pady=5, padx=10)
        cname_txt_XL.config(state=tk.DISABLED)

        self.rtfPath.trace_add("write", self.verifier_chemin_fichiers)

        open_button_XL = tk.Button(
            F2,
            text='Parcourir',
            command=self.select_files_XL
        )
        open_button_XL.grid(row=0, column=2, pady=5, padx=10)        

        # Frame pour le lancement du traitement
        F3 = tk.LabelFrame(self.root, text="Vérification", font=('times new roman', 14, 'bold'), bd=2, fg="Black", bg="#6fa8dc")
        # F3.place(x=0, y=300, relwidth=1, height=300)
        F3.place(x=0, y=240, relwidth=1, relheight=0.6)
        execute_lbl = tk.Label(F3, text="Lancer la vérification :", bg=bg_color, font=('times new roman', 15, 'normal'))
        execute_lbl.grid(row=0, column=0, padx=10, pady=5, sticky="e")
        self.execute_button = tk.Button(
            F3,
            text='Démarrer',
            command=self.process_exigences
        )
        self.execute_button.grid(row=0, column=1, pady=5, padx=10, sticky="w")        

        # text_area = scrolledtext.ScrolledText(F3, wrap=tk.WORD ,width=*, height=50, font=("Times New Roman", 15)) 
        # self.text_area = scrolledtext.ScrolledText(F3, wrap=tk.WORD ,font=("Times New Roman", 15)) 
        
        # self.text_area.grid(row=1, padx=10,column=0, columnspan=2, sticky='w'+'e'+'n'+'s') 

        # scroll_y = tk.Scrollbar(F3, orient=tk.VERTICAL)
        # self.txtarea = tk.Text(F3, yscrollcommand=scroll_y.set)
        # scroll_y.pack(side=tk.RIGHT, fill=tk.Y)
        # scroll_y.config(command=self.txtarea.yview)
        # self.txtarea.pack(fill=tk.BOTH, expand=1)    


        self.text_area = tk.Text(F3, wrap="word", width=40, height=10)
        self.text_area.grid(row=1, column=0, columnspan=2, padx=10, pady=10, sticky="nsew")

        scroll_y = tk.Scrollbar(F3, orient="vertical", command=self.text_area.yview)
        scroll_y.grid(row=1, column=2, sticky="ns")

        self.text_area.config(yscrollcommand=scroll_y.set, exportselection=True, selectbackground="lightblue")

        # Configurer le gestionnaire de géométrie pour que la ligne et la colonne s'agrandissent
        F3.grid_rowconfigure(1, weight=1)
        F3.grid_columnconfigure(0, weight=1)
        F3.grid_columnconfigure(1, weight=1)

        #open_button.pack(expand=True)
        self.Init_Chemins()
        self.verifier_chemin_fichiers()
        pass


        
    def __createIniFileIfNotExist(self, config):
        if not os.path.exists(FICHIER_INI):
            # Ajoutez des sections et des options
            config.add_section("Parametres")
            config.set('Parametres', 'specsPath', '')
            config.set('Parametres', 'rtfPath', '')

            # Écrivez les données dans le fichier
            with open(FICHIER_INI, 'w') as fichier_ini:
                config.write(fichier_ini)            


    def lire_configuration(self):
        # Vérifier si la section "Parametres" existe
        if 'Parametres' in self.config:
            # Lire les informations spécifiques
            self.specsPathConfig = self.config.get('Parametres', 'specsPath', fallback='')
            self.rtfPathConfig = self.config.get('Parametres', 'rtfPath', fallback='')
        else:
            print("La section 'Parametres' n'a pas été trouvée dans le fichier INI.")
            self.specsPathConfig = ""
            self.rtfPathConfig = ""


    def sauvegarder_informations(self):
        self.config.set('Parametres', 'specsPath', self.specsPath.get())
        self.config.set('Parametres', 'rtfPath', self.rtfPath.get())
        # Écrire la configuration dans le fichier INI
        with open(FICHIER_INI, 'w') as configfile:
            self.config.write(configfile)        

    def Init_Chemins(self):
        self.specsPath.set(self.specsPathConfig)
        self.rtfPath.set(self.rtfPathConfig)

    def verifier_chemin_fichiers(self, *args):
        chemin_fichier_specs = self.specsPath.get()
        chemin_fichier_rtf = self.rtfPath.get()
        specs_ok = os.path.exists(chemin_fichier_specs)
        rtf_ok = os.path.exists(chemin_fichier_rtf)
        if specs_ok and rtf_ok:
            self.execute_button["state"] = "normal"  # Activer le bouton si le chemin existe
        else:
            self.execute_button["state"] = "disabled"  # Désactiver le bouton sinon


    def select_files(self):
        filetypes = (
            ('word file', '*.docx')
        )

        filename = fd.askopenfilename(
            title='Open a file',
            initialdir='/',
            filetypes=[("word files", "*.docx")])

        # messagebox.showinfo(
        #     title='Selected File',
        #     message=filename
        # )
        self.specsPath.set(filename)

    def select_files_XL(self):
        filetypes = [("excel files", "*.xlsx")]

        filename = fd.askopenfilename(
            title='Open a file',
            initialdir='/',
            filetypes=filetypes)

        # messagebox.showinfo(
        #     title='Selected File',
        #     message=filename
        # )
        self.rtfPath.set(filename)

    #commenter
    def process_exigences(self):
        word_file_path = self.specsPath.get()
        excel_file_path = self.rtfPath.get()
        self.text_area.delete(1.0, tk.END)
        self.text_area.insert(tk.END, f"==> Extraction des exigences du document word : {word_file_path}"+ '\n')
        extracted_reqs = self.__extract_requirements_from_word(word_file_path)
        ordered_Word_reqs = sorted(list(extracted_reqs))
        for req in ordered_Word_reqs:
            self.text_area.insert(tk.END, req + '\n')

        self.text_area.insert(tk.END,"============================================================"+ '\n')
        self.text_area.insert(tk.END,f"Nombre total d'exigences extraites du document word : {len(extracted_reqs)}"+ '\n')
        self.text_area.insert(tk.END,"============================================================"+ '\n')

        self.text_area.insert(tk.END, "" + '\n')
        self.text_area.insert(tk.END, f"==> Extraction des exigences du document Excel : {excel_file_path}"+ '\n')

        XL_extracted_reqs = self.__extract_requirements_from_excel(excel_file_path, "Détail des tests")
        ordered_XL_reqs = sorted(list(XL_extracted_reqs))
        for req in ordered_XL_reqs:
            self.text_area.insert(tk.END, req + '\n')

        self.text_area.insert(tk.END,"============================================================"+ '\n')
        self.text_area.insert(tk.END,f"Nombre total d'exigences extraites du document Excel : {len(XL_extracted_reqs)}"+ '\n')
        self.text_area.insert(tk.END,"============================================================"+ '\n')
        self.text_area.insert(tk.END,""+ '\n')

        # Liste des exigences présentes dans le document MS Word mais absentes dans le document Excel        
        missing_in_excel = set(extracted_reqs) - set(XL_extracted_reqs)

        # Affichage des exigences manquantes
        self.text_area.insert(tk.END,"============================================================"+ '\n')
        self.text_area.insert(tk.END,f"Exigences présentes dans le document MS Word mais absentes dans le document Excel: {len(missing_in_excel)}"+ '\n')
        self.text_area.insert(tk.END,"============================================================"+ '\n')

        self.text_area.insert(tk.END,""+ '\n')

        self.text_area.insert(tk.END,"==> Bilan : "+ '\n')
        nb_exigences_test = len(XL_extracted_reqs)
        nb_exigences = len(extracted_reqs)
        nb_exigences_non_test = len(missing_in_excel)
        pc_exigences_test = (nb_exigences_test*100)/nb_exigences
        pc_exigences_non_test = (nb_exigences_non_test*100)/nb_exigences
        self.text_area.insert(tk.END,f"  - Nb exigences total : {nb_exigences}"+ '\n')
        self.text_area.insert(tk.END,f"  - Nb exigences testées : {nb_exigences_test} => {pc_exigences_test:.2f}%"+ '\n')
        self.text_area.insert(tk.END,f"  - Nb exigences NON testées : {nb_exigences_non_test} => {pc_exigences_non_test:.2f}%"+ '\n')
        self.text_area.insert(tk.END,"" + '\n')
        self.text_area.insert(tk.END,"  - Exigences non testées : " + '\n')
        self.text_area.insert(tk.END,"    ----------------------- " + '\n')
        ordered_missing = sorted(list(missing_in_excel))
        for requirement in ordered_missing:
            self.text_area.insert(tk.END,requirement + '\n')

        # Liste des exigences présentes dans le document Excel mais absentes dans le document MS Word  
        missing_in_word = set(XL_extracted_reqs) - set(extracted_reqs)
        # Affichage des exigences manquantes
        self.text_area.insert(tk.END,""+ '\n')
        self.text_area.insert(tk.END,"============================================================"+ '\n')
        self.text_area.insert(tk.END,f"==> Nombre d'exigences présentes dans le document Excel mais absentes dans le document MS Word : {len(missing_in_word)}"+ '\n')
        self.text_area.insert(tk.END,"============================================================"+ '\n')
        for requirement in missing_in_word:
            print(requirement)            


    # Fonction pour extraire les numéros d'exigences depuis le document Word
    def __extract_requirements_from_word(self, doc_path):
        doc = Document(doc_path)
    #    print(doc)
        requirements = set()
        pattern_ALL = re.compile(r'[A-Z]+-[a-zA-ZÀ-ÖØ-öø-ÿ]+-\d{3}')
        for paragraph in doc.paragraphs:
            # Ajoutez ici la logique pour extraire les numéros d'exigences selon votre pattern
            matches_ALL = pattern_ALL.findall(paragraph.text)
            requirements.update(matches_ALL)

        # Traitement des tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches_ALL = pattern_ALL.findall(cell.text)
                    requirements.update(matches_ALL)

        return requirements

    def __extract_requirements_from_excel(self, excel_path, sheet_name):
        requirements_set = set()

        wb = openpyxl.load_workbook(excel_path)
        
        # Sélectionne l'onglet spécifié par sheet_name
        sheet = wb[sheet_name]

        # Obtient l'indice de la colonne spécifiée par column_name
        column_index = openpyxl.utils.column_index_from_string("L")
    

        # Expression régulière pour rechercher des motifs du type 'EF-xxx-***'
        pattern_ALL = re.compile(r'[A-Z]+-[a-zA-ZÀ-ÖØ-öø-ÿ]+-\d{3}')

        # Traitement des cellules dans la colonne spécifiée
        for column in sheet.iter_cols(min_col=column_index, max_col=column_index, min_row=2, max_row=sheet.max_row):
            for cell in column:
                cell_value = cell.value
                if cell_value and not str(cell_value).isspace():
                    matches_ALL = pattern_ALL.findall(str(cell_value))        
                    requirements_set.update(matches_ALL)

        return list(requirements_set)

if __name__ == "__main__" :
    root = tk.Tk()
    app = Exigences_App(root)
    root.mainloop()
    app.sauvegarder_informations()